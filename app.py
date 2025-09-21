from flask import Flask, render_template, request, jsonify, send_from_directory
import os
from datetime import datetime
from dotenv import load_dotenv
from werkzeug.utils import secure_filename
from s3_storage import S3Storage
from bedrock_provider import BedrockProvider
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
try:
    import docx
except ImportError:
    docx = None

# Load environment variables
load_dotenv()

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max

# Create uploads directory
os.makedirs('uploads', exist_ok=True)

class AILearningPlatform:
    def __init__(self):
        self.learning_formats = ['chat', 'visual', 'ebook']
        self.bedrock_available = False
        self.ai_provider = BedrockProvider()
        self.s3_storage = S3Storage()
        
        # Initialize AWS Transcribe
        try:
            import boto3
            # Use ap-southeast-1 to match S3 bucket region
            transcribe_region = 'ap-southeast-1'
            self.transcribe_client = boto3.client(
                'transcribe',
                aws_access_key_id=os.getenv('AWS_ACCESS_KEY_ID'),
                aws_secret_access_key=os.getenv('AWS_SECRET_ACCESS_KEY'),
                region_name=transcribe_region
            )
            print(f"Transcribe client initialized in region: {transcribe_region}")
            print("AWS Transcribe client initialized")
        except Exception as e:
            print(f"Failed to initialize Transcribe: {e}")
            self.transcribe_client = None
        
        # Using AWS Bedrock Nova Pro with full S3 storage
        print("AI Learning Platform initialized with AWS Bedrock Nova Pro and full S3 cloud storage")
        
    def extract_text_from_file(self, filepath):
        """Extract text from uploaded files"""
        try:
            if filepath.endswith('.txt'):
                with open(filepath, 'r', encoding='utf-8') as f:
                    content = f.read()[:15000]  # Increased from 5000 to 15000
                    print(f"TXT: Read {len(content)} chars")
                    return content
            elif filepath.endswith('.pdf') and PyPDF2:
                with open(filepath, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ""
                    total_pages = len(reader.pages)
                    
                    # Read up to 50 pages for comprehensive coverage
                    if total_pages > 50:
                        # For very large PDFs, read first 25 pages + every 3rd page after
                        key_pages = list(range(0, min(25, total_pages)))  # First 25 pages
                        key_pages.extend(range(25, total_pages, 3))  # Every 3rd page after
                        pages_to_read = key_pages[:50]  # Max 50 pages total
                    else:
                        pages_to_read = range(total_pages)  # Read all pages if 50 or less
                    
                    for page_num in pages_to_read:
                        if page_num < total_pages:
                            page_text = reader.pages[page_num].extract_text()
                            text += f"[Page {page_num + 1}] {page_text}\n\n"
                            
                            # Stop if we have enough content
                            if len(text) > 50000:  # Increased for 50-page coverage
                                break
                    
                    print(f"PDF: Read {len(pages_to_read)} key pages from {total_pages} total, {len(text)} chars")
                    return text[:50000]  # Return up to 50k characters
            elif filepath.endswith('.docx') and docx:
                doc = docx.Document(filepath)
                text = ""
                for para in doc.paragraphs[:50]:  # Increased from 20 to 50 paragraphs
                    text += para.text + "\n"
                    if len(text) > 15000:  # Stop if we have enough
                        break
                print(f"DOCX: Read {len([p for p in doc.paragraphs if p.text.strip()])} paragraphs, {len(text)} chars")
                return text[:15000]
            return ""
        except Exception as e:
            print(f"Error extracting text: {e}")
            return ""
    
    def get_bedrock_response(self, topic, complexity_level, format_type, context=""):
        """Get AI response from AWS Bedrock"""
        
        # Use AI provider instead of fallback
        return self.ai_provider.get_ai_response(topic, complexity_level, format_type, context)
        
        level_prompts = {
            'beginner': 'Explain in very simple terms that anyone can understand',
            'intermediate': 'Explain with moderate detail and some technical terms',
            'advanced': 'Provide comprehensive explanation with technical depth'
        }
        
        format_prompts = {
            'chat': 'in a conversational, friendly tone with questions and interactive elements',
            'sketch': 'with detailed visual descriptions, ASCII diagrams, and step-by-step visual breakdowns',
            'video': 'as a complete video script with timestamps, scene descriptions, and narration cues',
            'ebook': 'as structured chapters with headings, bullet points, examples, and exercises'
        }
        
        context_prompt = f" Reference this additional context: {context}" if context else ""
        prompt = f"You are an expert educator. {level_prompts[complexity_level]} the topic '{topic}' with specific facts, examples, and details.{context_prompt} {format_prompts[format_type]}. Include real names, dates, companies, or specific examples related to {topic}. Avoid generic explanations - be topic-specific with concrete information."
        
        try:
            body = {
                "anthropic_version": "bedrock-2023-05-31",
                "max_tokens": 1000,
                "messages": [{"role": "user", "content": prompt}]
            }
            
            response = self.bedrock.invoke_model(
                modelId='anthropic.claude-3-haiku-20240307-v1:0',
                body=json.dumps(body)
            )
            
            response_body = json.loads(response['body'].read())
            return response_body['content'][0]['text'].strip()
            
        except Exception as e:
            print(f"Bedrock error: {e}")
            return self.get_smart_fallback(topic, complexity_level, format_type)
    
    def get_smart_fallback(self, topic, complexity_level, format_type):
        """Smart fallback system with good explanations"""
        
        # Skip S3 database check for now - go straight to AI
        print(f"🚀 Skipping S3 database, using AI for {topic}")
        # topic_data = self.s3_db.get_topic(topic)
        # if topic_data:
        #     return topic_data['explanation']
        
        # Generic fallback responses for when AI fails
        topic_lower = topic.lower()
        if complexity_level == 'primary':
            return f"{topic.capitalize()} is something really cool that we can learn about! It's like when you discover how things work around you, and it helps us understand our world better."
        elif complexity_level == 'secondary':
            return f"{topic.capitalize()} is an important concept that connects to many things we study. It involves understanding key principles and how they work together in real situations."
        elif complexity_level == 'foundation':
            return f"{topic.capitalize()} involves multiple principles and concepts that work together, with practical applications and real-world implications that are relevant to your studies."
        else:  # degree
            return f"{topic.capitalize()} requires comprehensive analysis, theoretical understanding, and synthesis of complex interconnected concepts and methodologies within its field of study."
        
        # Comprehensive topic-specific explanations
        explanations = {
            'cryptocurrency': {
                'beginner': "Cryptocurrency is digital money stored on computers. Like Bitcoin, you can send it directly to others without using banks. It uses special math codes to keep transactions secure and prevent counterfeiting.",
                'intermediate': "Cryptocurrency operates on blockchain networks using cryptographic algorithms. Miners validate transactions through proof-of-work consensus, creating immutable ledgers that eliminate intermediaries while maintaining decentralized control.",
                'advanced': "Cryptocurrency leverages distributed ledger technology with cryptographic hash functions, merkle trees, and consensus mechanisms like PoW/PoS to achieve Byzantine fault tolerance in trustless peer-to-peer value transfer systems."
            },
            'machine learning': {
                'beginner': "Machine learning teaches computers to recognize patterns by showing them lots of examples. Like teaching a child to recognize cats by showing many cat pictures, computers learn to make predictions.",
                'intermediate': "Machine learning uses statistical algorithms to find patterns in data. Neural networks, decision trees, and regression models train on datasets to make predictions without explicit programming for each scenario.",
                'advanced': "Machine learning employs gradient descent optimization, backpropagation, and regularization techniques across supervised, unsupervised, and reinforcement learning paradigms to minimize loss functions and generalize from training data."
            },
            'python': {
                'beginner': "Python is a programming language that's easy to read and write. It uses simple English-like commands to tell computers what to do, making it perfect for beginners to learn coding.",
                'intermediate': "Python is an interpreted, high-level language with dynamic typing and automatic memory management. Its extensive standard library and frameworks like Django make it versatile for web development, data science, and automation.",
                'advanced': "Python implements duck typing with a global interpreter lock (GIL), uses reference counting with cycle detection for garbage collection, and supports metaclasses, decorators, and context managers for advanced programming patterns."
            },
            'quantum computing': {
                'beginner': "Quantum computers use tiny particles that can be in multiple states at once, unlike regular computers that use just 0s and 1s. This lets them solve certain problems much faster.",
                'intermediate': "Quantum computing exploits quantum superposition and entanglement to process information. Qubits can exist in multiple states simultaneously, enabling parallel computation through quantum algorithms like Shor's and Grover's.",
                'advanced': "Quantum computing utilizes quantum mechanical phenomena including superposition, entanglement, and quantum interference. Gate-based quantum circuits manipulate qubit states through unitary operations to achieve quantum speedup for specific computational problems."
            },
            'blockchain': {
                'beginner': "Blockchain is like a digital ledger that everyone can see but no one can cheat. Each new transaction gets added as a 'block' and linked to previous blocks, creating an unchangeable chain.",
                'intermediate': "Blockchain creates immutable distributed ledgers through cryptographic hashing and consensus mechanisms. Each block contains transaction data, timestamps, and hash pointers, forming a tamper-evident chain validated by network participants.",
                'advanced': "Blockchain implements cryptographic hash functions, merkle trees, and distributed consensus protocols (PoW, PoS, PBFT) to achieve Byzantine fault tolerance in decentralized systems while maintaining data integrity and preventing double-spending attacks."
            }
        }
        
        # Check for specific topics with partial matching
        topic_lower = topic.lower()
        for key, levels in explanations.items():
            if key in topic_lower or any(word in topic_lower for word in key.split()):
                return levels.get(complexity_level, levels['beginner'])
        
        # Enhanced dynamic responses based on topic keywords
        tech_keywords = ['ai', 'artificial intelligence', 'neural', 'algorithm', 'data', 'computer', 'software', 'programming']
        science_keywords = ['physics', 'chemistry', 'biology', 'mathematics', 'theory', 'scientific']
        business_keywords = ['marketing', 'finance', 'economics', 'business', 'management', 'strategy']
        
        if any(keyword in topic_lower for keyword in tech_keywords):
            if complexity_level == 'primary':
                return f"{topic.capitalize()} is like a smart helper that uses computers to solve problems! It's like having a robot friend that can think and help us do cool things."
            elif complexity_level == 'secondary':
                return f"{topic.capitalize()} is a technology concept that helps solve problems using computers and data. It involves step-by-step processes and logical thinking to create solutions."
            elif complexity_level == 'foundation':
                return f"{topic.capitalize()} combines technical principles with practical applications, using structured approaches to process information and solve real-world problems."
            else:  # degree
                return f"{topic.capitalize()} involves complex computational algorithms, data structures, and systematic methodologies requiring deep technical understanding and implementation expertise."
        
        elif any(keyword in topic_lower for keyword in science_keywords):
            if complexity_level == 'primary':
                return f"{topic.capitalize()} is about how nature works! It's like being a detective and discovering cool secrets about the world around us."
            elif complexity_level == 'secondary':
                return f"{topic.capitalize()} is a scientific concept that explains how things work in nature. It helps us understand patterns and make predictions about the world around us."
            elif complexity_level == 'foundation':
                return f"{topic.capitalize()} involves scientific principles and experimental methods to understand natural phenomena and their underlying mechanisms."
            else:  # degree
                return f"{topic.capitalize()} encompasses rigorous scientific methodologies, mathematical modeling, and empirical research requiring analytical reasoning and theoretical frameworks."
        
        elif any(keyword in topic_lower for keyword in business_keywords):
            if complexity_level == 'primary':
                return f"{topic.capitalize()} is about how people work together to make and sell things! It's like running a lemonade stand but bigger."
            elif complexity_level == 'secondary':
                return f"{topic.capitalize()} is about how organizations work and make decisions. It involves understanding people, money, and strategies to achieve goals."
            elif complexity_level == 'foundation':
                return f"{topic.capitalize()} involves business principles, market analysis, and organizational strategies to create value and achieve objectives."
            else:  # degree
                return f"{topic.capitalize()} requires strategic analysis, market dynamics understanding, and complex decision-making frameworks within competitive business environments."
        
        # Generic fallback
        if complexity_level == 'beginner':
            return f"{topic.capitalize()} is an important concept with practical applications. It involves understanding key principles and how they work together in real situations."
        elif complexity_level == 'advanced':
            return f"{topic.capitalize()} requires comprehensive analysis, theoretical understanding, and synthesis of complex interconnected concepts and methodologies."
        else:
            return f"{topic.capitalize()} involves multiple principles and concepts that work together, with practical applications and real-world implications."
        
    def simplify_topic(self, topic, complexity_level, format_type, uploaded_files=None):
        """Main function to get AI explanation for any topic"""
        print(f"📚 Simplifying topic: {topic} ({complexity_level}, {format_type})")
        
        context = ""
        if uploaded_files:
            for filepath in uploaded_files:
                extracted = self.extract_text_from_file(filepath)
                context += extracted + "\n"
                print(f"Extracted {len(extracted)} chars from {filepath}")
        
        print(f"Total context: {len(context)} chars")
        
        # Use AI for all content generation including video
        print(f"Getting AI response for {topic}")
        ai_content = self.ai_provider.get_ai_response(topic, complexity_level, format_type, context)
        
        if not ai_content:
            print(f"AI returned empty, using fallback for {topic}")
            ai_content = self.get_smart_fallback(topic, complexity_level, format_type)
        else:
            print(f"AI response received: {len(ai_content)} chars")
                
        print(f"🎨 Formatting content for {format_type}")
        formatted_content = self.format_content(ai_content, format_type)
        
        return {
            'content': formatted_content,
            'format': format_type,
            'level': complexity_level
        }
    
    def format_content(self, content, format_type):
        """Apply format-specific styling to content"""
        if format_type == 'chat':
            return f"{content}\n\n💬 Ready to dive deeper? What specific aspect interests you most?"
        elif format_type in ['sketch', 'visual']:
            return content

        elif format_type == 'ebook':
            return f"<h2>📚 Document Summary</h2>\n\n{content}"
        
        return content
    
    def transcribe_audio(self, audio_url, filename):
        """Transcribe audio using AWS Transcribe"""
        if not self.transcribe_client:
            return None
            
        try:
            job_name = f"transcribe_{filename.replace('.', '_')}_{int(datetime.now().timestamp())}"
            
            # Start transcription job
            print(f"Starting transcription job: {job_name}")
            print(f"Audio URL: {audio_url}")
            
            self.transcribe_client.start_transcription_job(
                TranscriptionJobName=job_name,
                Media={'MediaFileUri': audio_url},
                MediaFormat='wav',
                LanguageCode='en-US'
            )
            print("Transcription job started successfully")
            
            # Wait for completion (polling)
            import time
            max_wait = 60  # Maximum 60 seconds
            wait_time = 0
            
            while wait_time < max_wait:
                response = self.transcribe_client.get_transcription_job(
                    TranscriptionJobName=job_name
                )
                
                status = response['TranscriptionJob']['TranscriptionJobStatus']
                
                if status == 'COMPLETED':
                    # Get transcript
                    transcript_uri = response['TranscriptionJob']['Transcript']['TranscriptFileUri']
                    
                    # Download and parse transcript
                    import requests
                    import json
                    
                    transcript_response = requests.get(transcript_uri)
                    transcript_data = transcript_response.json()
                    
                    transcript_text = transcript_data['results']['transcripts'][0]['transcript']
                    
                    # Clean up job
                    try:
                        self.transcribe_client.delete_transcription_job(
                            TranscriptionJobName=job_name
                        )
                    except:
                        pass
                    
                    return transcript_text
                    
                elif status == 'FAILED':
                    failure_reason = response['TranscriptionJob'].get('FailureReason', 'Unknown error')
                    print(f"Transcription failed: {failure_reason}")
                    return f"Transcription failed: {failure_reason}"
                    
                time.sleep(2)
                wait_time += 2
            
            # Cleanup if job didn't complete
            try:
                self.transcribe_client.delete_transcription_job(
                    TranscriptionJobName=job_name
                )
            except:
                pass
                
            return None
            
        except Exception as e:
            print(f"Transcription error: {e}")
            import traceback
            traceback.print_exc()
            return f"Error: {str(e)}"

platform = AILearningPlatform()

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/learn', methods=['POST'])
def learn():
    try:
        topic = request.form.get('topic')
        level = request.form.get('level', 'primary')
        format_type = request.form.get('format', 'chat')
        
        print(f"📝 Learn request: topic='{topic}', level='{level}', format='{format_type}'")
        
        if not topic:
            return jsonify({'error': 'Topic is required'}), 400
        
        uploaded_files = []
        if 'documents' in request.files:
            files = request.files.getlist('documents')
            for file in files:
                if file.filename:
                    filename = secure_filename(file.filename)
                    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    file.save(filepath)
                    uploaded_files.append(filepath)
        
        result = platform.simplify_topic(topic, level, format_type, uploaded_files)
        print(f"✅ Learn result: {len(str(result))} chars")
        return jsonify(result)
    except Exception as e:
        print(f"ERROR in learn route: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Something went wrong, please try again. Debug: {str(e)}'}), 500

@app.route('/formats')
def get_formats():
    return jsonify(platform.learning_formats)

@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/forum/threads', methods=['GET'])
def get_forum_threads():
    try:
        threads = platform.s3_storage.get_forum_threads()
        return jsonify({'success': True, 'threads': threads})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/forum/create_thread', methods=['POST'])
def create_forum_thread():
    try:
        data = request.get_json()
        result = platform.s3_storage.create_forum_thread(
            data.get('user_id'),
            data.get('title'),
            data.get('content'),
            data.get('topic'),
            data.get('level')
        )
        return jsonify(result)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/forum/like_thread', methods=['POST'])
def like_forum_thread():
    try:
        data = request.get_json()
        result = platform.s3_storage.like_forum_thread(
            data.get('user_id'),
            data.get('thread_id')
        )
        return jsonify(result)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/leaderboard/<filter_type>', methods=['GET'])
def get_leaderboard(filter_type):
    try:
        leaderboard = platform.s3_storage.get_knowledge_leaderboard(filter_type)
        return jsonify({'success': True, 'leaderboard': leaderboard})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/analyze_blurting', methods=['POST'])
def analyze_blurting():
    try:
        data = request.get_json()
        topic = data.get('topic')
        level = data.get('level')
        blurt_text = data.get('blurt_text')
        time_spent = data.get('time_spent', 0)
        
        if not topic or not blurt_text:
            return jsonify({'success': False, 'error': 'Topic and blurt text are required'})
        
        # Analyze blurting using AI
        analysis = platform.ai_provider.analyze_blurting(topic, level, blurt_text, time_spent)
        
        if analysis:
            return jsonify({'success': True, 'analysis': analysis})
        else:
            return jsonify({'success': False, 'error': 'Failed to analyze blurting session'})
            
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/generate_flashcards', methods=['POST'])
def generate_flashcards():
    try:
        data = request.get_json()
        topic = data.get('topic')
        level = data.get('level')
        
        if not topic:
            return jsonify({'success': False, 'error': 'Topic is required'})
        
        # Generate flashcards using AI
        flashcards = platform.ai_provider.generate_flashcards(topic, level)
        
        if flashcards:
            return jsonify({'success': True, 'flashcards': flashcards})
        else:
            return jsonify({'success': False, 'error': 'Failed to generate flashcards'})
            
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/transcribe_audio', methods=['POST'])
def transcribe_audio():
    try:
        if 'audio' not in request.files:
            return jsonify({'success': False, 'error': 'No audio file provided'})
        
        audio_file = request.files['audio']
        if audio_file.filename == '':
            return jsonify({'success': False, 'error': 'No audio file selected'})
        
        # Save audio file temporarily
        filename = secure_filename(f"audio_{datetime.now().strftime('%Y%m%d_%H%M%S')}.wav")
        audio_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        audio_file.save(audio_path)
        
        # Upload to S3 for Transcribe
        s3_key = f"audio/{filename}"
        try:
            platform.s3_storage.s3_client.upload_file(
                audio_path, 
                platform.s3_storage.bucket_name, 
                s3_key
            )
            audio_url = f"s3://{platform.s3_storage.bucket_name}/{s3_key}"
        except Exception as e:
            return jsonify({'success': False, 'error': f'S3 upload failed: {str(e)}'})
        

        
        # Use AWS Transcribe
        transcript = platform.transcribe_audio(audio_url, filename)
        
        # Clean up local file
        os.remove(audio_path)
        
        if transcript:
            return jsonify({'success': True, 'transcript': transcript})
        else:
            return jsonify({'success': False, 'error': 'Transcription failed'})
            
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/submit_explanation', methods=['POST'])
def submit_explanation():
    try:
        user_id = request.form.get('user_id', 'anonymous')
        topic = request.form.get('topic')
        level = request.form.get('level')
        transcript = request.form.get('transcript')
        
        # Get AI grading from Bedrock
        ai_score = platform.ai_provider.grade_explanation(topic, level, transcript)
        
        # Submit to S3 with AI score
        result = platform.s3_storage.submit_explanation(user_id, topic, level, transcript, ai_score)
        return jsonify(result)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/community_explanations/<topic>')
def get_community_explanations(topic):
    try:
        explanations = platform.s3_storage.get_community_explanations(topic)
        return jsonify({'explanations': explanations})
    except Exception as e:
        return jsonify({'error': str(e)})

@app.route('/upvote_explanation', methods=['POST'])
def upvote_explanation():
    try:
        user_id = request.form.get('user_id', 'anonymous')
        explanation_id = request.form.get('explanation_id')
        
        result = platform.s3_storage.upvote_explanation(user_id, explanation_id)
        return jsonify(result)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/user_stats/<user_id>')
def get_user_stats(user_id):
    try:
        stats = platform.s3_storage.get_user_stats(user_id)
        return jsonify(stats)
    except Exception as e:
        return jsonify({'error': str(e)})

# Game System Routes
@app.route('/game/challenge/<difficulty>')
@app.route('/game/challenge/<difficulty>/<category>')
def get_challenge(difficulty, category=None):
    try:
        challenge = platform.s3_storage.get_random_challenge(difficulty, category)
        if challenge:
            # Don't send correct answer to frontend
            challenge.pop('correct_answer', None)
            return jsonify(challenge)
        return jsonify({'error': 'No challenges available'})
    except Exception as e:
        return jsonify({'error': str(e)})

@app.route('/game/categories')
def get_categories():
    try:
        categories = platform.s3_storage.get_categories()
        return jsonify({'categories': categories})
    except Exception as e:
        return jsonify({'error': str(e)})

@app.route('/game/submit', methods=['POST'])
def submit_challenge():
    try:
        user_id = request.form.get('user_id')
        challenge_id = request.form.get('challenge_id')
        answer = request.form.get('answer')
        time_taken = int(request.form.get('time_taken', 0))
        
        result = platform.s3_storage.submit_challenge_answer(user_id, challenge_id, answer, time_taken)
        return jsonify(result)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/game/stats/<user_id>')
def get_game_stats(user_id):
    try:
        stats = platform.s3_storage.get_player_stats(user_id)
        return jsonify(stats)
    except Exception as e:
        return jsonify({'error': str(e)})

@app.route('/game/leaderboard')
def get_game_leaderboard():
    try:
        leaderboard = platform.s3_storage.get_leaderboard()
        return jsonify({'leaderboard': leaderboard})
    except Exception as e:
        return jsonify({'error': str(e)})

if __name__ == '__main__':
    app.run(debug=True)